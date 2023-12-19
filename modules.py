import sys
import io
import xlsxwriter
from docx import Document
from pptx import Presentation
from docx.shared import Inches, Mm
from random import randint, choice
from datetime import datetime, time, timedelta

from PyQt5 import uic  # Импортируем uic
from PyQt5.QtWidgets import QApplication, QMainWindow, QWidget, QInputDialog, QFileDialog, QPushButton


class AddHall(QMainWindow):
    def __init__(self, parent):
        super().__init__()
        uic.loadUi("addHall.ui", self)
        self.parent = parent

        for c in self.parent.cinemas:
            self.cinemaInput.addItem(c)

        self.addBtn.clicked.connect(self.add)

    def add(self):
        cinema = self.cinemaInput.currentText()
        column = self.columnInput.value()
        row = self.rowInput.value()

        self.parent.cinemas[cinema].append([column, row])
        self.parent.resultLabel.setText('Зал добавлен')
        self.close()


class SetChairs(QMainWindow):
    def __init__(self, parent):
        super().__init__()
        uic.loadUi("setChairs.ui", self)
        self.parent = parent

        self.cinemas = []
        for c in self.parent.cinemas:
            if len(self.parent.cinemas[c]):
                self.cinemaInput.addItem(c)

        self.change_halls()

        self.cinemaInput.currentTextChanged.connect(self.change_halls)
        self.saveBtn.clicked.connect(self.save)

    def change_halls(self):
        self.hallInput.clear()
        for i in range(len(self.parent.cinemas[self.cinemaInput.currentText()])):
            self.hallInput.addItem(str(i + 1))

    def save(self):
        cinema = self.cinemaInput.currentText()
        hall = self.hallInput.currentIndex()
        column = self.columnInput.value()
        row = self.rowInput.value()

        self.parent.cinemas[cinema][hall] = [column, row]
        self.parent.resultLabel.setText('Изменение сохранено')
        self.close()


class CreateSession(QMainWindow):
    def __init__(self, parent):
        super().__init__()
        uic.loadUi("createSession.ui", self)
        self.parent = parent

        self.cinemas = []
        for c in self.parent.cinemas:
            if len(self.parent.cinemas[c]):
                self.cinemaInput.addItem(c)

        self.change_halls()

        self.dateInput.setDate(datetime.now())
        self.startInput.setTime(datetime.now().time())
        self.durationInput.setTime(time(hour=2))

        self.cinemaInput.currentTextChanged.connect(self.change_halls)
        self.createBtn.clicked.connect(self.create)

    def change_halls(self):
        self.hallInput.clear()
        for i in range(len(self.parent.cinemas[self.cinemaInput.currentText()])):
            self.hallInput.addItem(str(i + 1))

    def create(self):
        cinema = self.cinemaInput.currentText()
        hall = self.hallInput.currentIndex()
        date = self.dateInput.date()
        start = self.startInput.time()
        duration = self.durationInput.time()
        film = self.filmInput.text()

        end = start.addSecs(duration.hour() * 3600 + duration.minute() * 60)

        for s in self.parent.sessions:
            if cinema == s['cinema'] and cinema == s['cinema'] and s['date'] == date:
                # end = s['start'].addSecs(s['duration'].hour() * 3600 + s['duration'].minute() * 60)
                if s['hall_num'] == hall + 1 and (start <= s['start'] < end
                                                  or start < s['start'].addSecs(
                            s['duration'].hour() * 3600 + s['duration'].minute() * 60) <= end):
                    self.statusBar().showMessage(f'В это время идёт сеанс: {session_in_text(s)}')
                    return
        else:
            column, row = self.parent.cinemas[cinema][hall]
            hall_num = hall + 1
            hall = [[0] * row for i in range(column)]
            session = {'cinema': cinema, 'hall': hall, 'date': date, 'start': start,
                       'duration': duration, 'film': film, 'hall_num': hall_num}
            self.parent.sessions.append(session)
            self.parent.resultLabel.setText(f'Добавлен сеанс: "{session_in_text(session)}"')
            self.close()


class SellTicket(QMainWindow):
    def __init__(self, parent):
        super().__init__()
        uic.loadUi("sellTicket.ui", self)
        self.parent = parent

        for s in self.parent.sessions:
            self.sessionInput.addItem(session_in_text(s))

        self.change_halls()

        self.sessionInput.currentTextChanged.connect(self.change_halls)
        self.sellBtn.clicked.connect(self.sell)

    def change_halls(self):
        len1 = len(self.parent.sessions[self.sessionInput.currentIndex()]['hall'])
        len2 = len(self.parent.sessions[self.sessionInput.currentIndex()]['hall'][0])
        self.columnInput.setMaximum(len1)
        self.rowInput.setMaximum(len2)

    def sell(self):
        session = self.sessionInput.currentIndex()
        column = self.columnInput.value() - 1
        row = self.rowInput.value() - 1
        if self.parent.sessions[session]['hall'][column][row]:
            self.statusBar().showMessage("Это место уже занято", 5000)
        else:
            self.parent.sessions[session]['hall'][column][row] = 1
            self.parent.resultLabel.setText('Билет продан')
            self.close()


class HallPlan(QWidget):
    def __init__(self, hall):
        super().__init__()
        self.setGeometry(100, 100, len(hall) * 70, len(hall[0]) * 50)
        self.setFixedSize(len(hall) * 70 + 20, len(hall[0]) * 50 + 20)
        self.btns = []
        for i in range(len(hall)):
            self.btns.append([])
            for j in range(len(hall[0])):
                self.btns[i].append(QPushButton(self))
                self.btns[i][j].setText(f"Ряд: {i + 1}\n Место: {j + 1}")
                self.btns[i][j].move(70 * j + 10, 50 * i + 10)
                if hall[i][j]:
                    self.btns[i][j].setStyleSheet(
                        'QPushButton {background-color: #FF0000; border: none; width: 65px; height: 45px;}')
                else:
                    self.btns[i][j].setStyleSheet(
                        'QPushButton {background-color: #00FF00; border: none; width: 65px; height: 45px;}')


class SeatsInRow(QMainWindow):
    def __init__(self, parent):
        super().__init__()
        uic.loadUi("seatsInRow.ui", self)
        self.parent = parent

        for c in self.parent.cinemas:
            self.cinemaInput.addItem(c)

        self.searchBtn.clicked.connect(self.search)

    def search(self):
        cinema = self.cinemaInput.currentText()
        n = self.seatsInput.value()

        closest = None
        for s in self.parent.sessions:
            if (s['cinema'] == cinema and s['date'] >= datetime.now().date()
                    and any(['0' * n in ''.join(map(str, row)) for row in s['hall']])):
                if not closest:
                    closest = s
                elif s['date'] < closest['date']:
                    closest = s
                elif s['date'] == closest['date'] and s['start'] < closest['start']:
                    closest = s

        if closest:
            start = f'{str(closest["start"].hour()).rjust(2, "0")}:{str(closest["start"].minute()).rjust(2, "0")}'
            self.parent.resultLabel.setText(
                f'Ближайший сеанс на который есть {n} мест в ряд\n'
                f'пройдёт {closest["date"].toString("dd.MM.yyyy")} в {start}.')
        else:
            self.parent.resultLabel.setText('Такой возможности нет')
        self.close()


class AdBooklet(QMainWindow):
    def __init__(self, parent):
        super().__init__()
        uic.loadUi("adBooklet.ui", self)
        self.parent = parent

        self.cinemas = []
        self.fname = None

        self.dateInput.setDate(datetime.now().date())

        self.doneBtn.clicked.connect(self.create)
        self.addCinema.clicked.connect(self.add_cinema)
        self.chooseImg.clicked.connect(self.choose_img)

    def create(self):
        film = self.filmInput.text()
        date = self.dateInput.date()
        description = self.descriptionInput.text()

        document = Document()

        heading = document.add_heading(film, level=0)
        heading.alignment = 1

        document.add_paragraph(description, style='Intense Quote')

        document.add_paragraph(f'В кино с {date.toString("dd.MM")}')

        document.add_paragraph('Только в кинотеатрах:')

        for c in self.cinemas:
            document.add_paragraph(c, style='List Bullet')

        if self.fname:
            p = document.add_paragraph()
            run = p.add_run()
            run.add_picture(self.fname, width=Mm(150))

        document.save(f'{film}.docx')
        self.parent.resultLabel.setText(f'Рекламный буклет здесь: {film}.docx')
        self.close()

    def add_cinema(self):
        cinema, ok_pressed = QInputDialog.getItem(self, 'Выберете кинотеатр',
                                                  "Выберете кинотеатр?",
                                                  set(self.parent.cinemas.keys()) - set(self.cinemas), 0,
                                                  False)
        if not ok_pressed:
            return

        self.cinemas.append(cinema)
        self.cinemasInput.setText(', '.join(self.cinemas))

    def choose_img(self):
        self.fname = QFileDialog.getOpenFileName(self, 'Выбрать картинку', '')[0]
        self.imagePath.setText(self.fname)
